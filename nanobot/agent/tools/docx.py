"""Word document operations tools using python-docx and docxtpl."""

import json
import re
from pathlib import Path
from typing import Any

from nanobot.agent.tools.base import Tool

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    from docx.shared import Pt, Inches, RGBColor, Twips
except ImportError as e:
    raise ImportError(
        "python-docx is required for Word document tools. "
        "Install it with: pip install nanobot-ai[docx]"
    ) from e


def _resolve_path(path: str, workspace: Path | None = None, allowed_dir: Path | None = None) -> Path:
    """Resolve path against workspace and enforce directory restriction."""
    p = Path(path).expanduser()
    if not p.is_absolute() and workspace:
        p = workspace / p
    resolved = p.resolve()
    if allowed_dir:
        try:
            resolved.relative_to(allowed_dir.resolve())
        except ValueError:
            raise PermissionError(f"Path {path} is outside allowed directory {allowed_dir}")
    return resolved


def _get_heading_level(paragraph) -> int | None:
    """Get heading level from paragraph, returns None if not a heading."""
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    if style_name.startswith("heading"):
        match = re.search(r"heading\s*(\d+)", style_name)
        if match:
            return int(match.group(1))
    if style_name.startswith("标题"):
        match = re.search(r"标题\s*(\d+)", style_name)
        if match:
            return int(match.group(1))
    return None


def _set_run_font(run, font_name: str | None, font_size: float | None, 
                  bold: bool | None, italic: bool | None, color: str | None) -> None:
    """Set font properties on a run."""
    if font_name:
        run.font.name = font_name
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), font_name)
    if font_size is not None:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


class DocxReadStructureTool(Tool):
    """Tool to read Word document structure."""

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "docx_read_structure"

    @property
    def description(self) -> str:
        return "Read the structure of a Word document, including paragraph count, table count, and heading hierarchy."

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The absolute path to the Word document (.docx)"
                }
            },
            "required": ["path"]
        }

    async def execute(self, path: str, **kwargs: Any) -> str:
        try:
            file_path = _resolve_path(path, self._workspace, self._allowed_dir)
            if not file_path.exists():
                return f"Error: File not found: {path}"
            if not file_path.suffix.lower() == ".docx":
                return f"Error: Not a .docx file: {path}"

            doc = Document(str(file_path))
            
            paragraphs_count = len(doc.paragraphs)
            tables_count = len(doc.tables)
            
            headings = []
            for i, para in enumerate(doc.paragraphs):
                level = _get_heading_level(para)
                if level:
                    headings.append({
                        "level": level,
                        "text": para.text.strip()[:100],
                        "index": i
                    })

            sections = [f"Section {i+1}" for i in range(len(doc.sections))] if doc.sections else []

            result = {
                "paragraphs_count": paragraphs_count,
                "tables_count": tables_count,
                "headings": headings,
                "sections_count": len(doc.sections),
                "sections": sections[:10]
            }

            return json.dumps(result, ensure_ascii=False, indent=2)

        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error reading document structure: {str(e)}"


class DocxReadTextTool(Tool):
    """Tool to read text content from Word document."""

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "docx_read_text"

    @property
    def description(self) -> str:
        return "Read all text content from a Word document, excluding tables."

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The absolute path to the Word document (.docx)"
                },
                "include_headings": {
                    "type": "boolean",
                    "description": "Whether to include heading text, default true"
                }
            },
            "required": ["path"]
        }

    async def execute(self, path: str, include_headings: bool = True, **kwargs: Any) -> str:
        try:
            file_path = _resolve_path(path, self._workspace, self._allowed_dir)
            if not file_path.exists():
                return f"Error: File not found: {path}"
            if not file_path.suffix.lower() == ".docx":
                return f"Error: Not a .docx file: {path}"

            doc = Document(str(file_path))
            
            text_parts = []
            for para in doc.paragraphs:
                level = _get_heading_level(para)
                if not include_headings and level:
                    continue
                if para.text.strip():
                    if level:
                        text_parts.append(f"[标题{level}] {para.text}")
                    else:
                        text_parts.append(para.text)

            return "\n\n".join(text_parts)

        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error reading document text: {str(e)}"


class DocxReadTablesTool(Tool):
    """Tool to read table data from Word document."""

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "docx_read_tables"

    @property
    def description(self) -> str:
        return "Read table data from a Word document, returns JSON format with all tables."

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The absolute path to the Word document (.docx)"
                },
                "table_index": {
                    "type": "integer",
                    "description": "Specific table index (0-based), omit to return all tables"
                }
            },
            "required": ["path"]
        }

    async def execute(self, path: str, table_index: int | None = None, **kwargs: Any) -> str:
        try:
            file_path = _resolve_path(path, self._workspace, self._allowed_dir)
            if not file_path.exists():
                return f"Error: File not found: {path}"
            if not file_path.suffix.lower() == ".docx":
                return f"Error: Not a .docx file: {path}"

            doc = Document(str(file_path))
            
            if not doc.tables:
                return json.dumps({"tables": [], "message": "No tables found in document"}, ensure_ascii=False)

            def table_to_data(table) -> list[list[str]]:
                data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip().replace("\n", " ")
                        row_data.append(cell_text)
                    data.append(row_data)
                return data

            if table_index is not None:
                if table_index < 0 or table_index >= len(doc.tables):
                    return f"Error: Table index {table_index} out of range (0-{len(doc.tables)-1})"
                
                table = doc.tables[table_index]
                result = {
                    "tables": [{
                        "index": table_index,
                        "rows": len(table.rows),
                        "cols": len(table.columns),
                        "data": table_to_data(table)
                    }]
                }
            else:
                tables_data = []
                for i, table in enumerate(doc.tables):
                    tables_data.append({
                        "index": i,
                        "rows": len(table.rows),
                        "cols": len(table.columns),
                        "data": table_to_data(table)
                    })
                result = {"tables": tables_data}

            return json.dumps(result, ensure_ascii=False, indent=2)

        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error reading tables: {str(e)}"


class DocxSetFontTool(Tool):
    """Tool to set font style in Word document."""

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "docx_set_font"

    @property
    def description(self) -> str:
        return "Set font style (font name, size, bold, italic, color) for specified content in a Word document."

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The absolute path to the Word document (.docx)"
                },
                "output_path": {
                    "type": "string",
                    "description": "Output path, omit to overwrite original file"
                },
                "target": {
                    "type": "string",
                    "enum": ["all", "paragraphs", "headings", "tables", "heading1", "heading2", "heading3"],
                    "description": "Target scope: all, paragraphs (body text), headings (all titles), tables, heading1/2/3 (specific level)"
                },
                "font_name": {
                    "type": "string",
                    "description": "Font name, e.g. 'Microsoft YaHei', 'SimSun', 'Arial'"
                },
                "font_size": {
                    "type": "number",
                    "description": "Font size in points, e.g. 12, 14, 16"
                },
                "bold": {
                    "type": "boolean",
                    "description": "Whether to make text bold"
                },
                "italic": {
                    "type": "boolean",
                    "description": "Whether to make text italic"
                },
                "color": {
                    "type": "string",
                    "description": "Font color in hex, e.g. 'FF0000' for red"
                }
            },
            "required": ["path", "target"]
        }

    async def execute(self, path: str, target: str, output_path: str | None = None,
                      font_name: str | None = None, font_size: float | None = None,
                      bold: bool | None = None, italic: bool | None = None,
                      color: str | None = None, **kwargs: Any) -> str:
        try:
            file_path = _resolve_path(path, self._workspace, self._allowed_dir)
            if not file_path.exists():
                return f"Error: File not found: {path}"
            if not file_path.suffix.lower() == ".docx":
                return f"Error: Not a .docx file: {path}"

            doc = Document(str(file_path))
            modified_count = 0

            def should_modify_paragraph(para) -> bool:
                if target == "all":
                    return True
                level = _get_heading_level(para)
                if target == "paragraphs":
                    return level is None
                if target == "headings":
                    return level is not None
                if target == "heading1":
                    return level == 1
                if target == "heading2":
                    return level == 2
                if target == "heading3":
                    return level == 3
                return False

            for para in doc.paragraphs:
                if should_modify_paragraph(para):
                    for run in para.runs:
                        _set_run_font(run, font_name, font_size, bold, italic, color)
                        modified_count += 1

            if target in ("all", "tables"):
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    _set_run_font(run, font_name, font_size, bold, italic, color)
                                    modified_count += 1

            save_path = _resolve_path(output_path, self._workspace, self._allowed_dir) if output_path else file_path
            save_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(save_path))

            return f"Successfully modified {modified_count} text runs. Saved to: {save_path}"

        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error setting font: {str(e)}"


class DocxSetParagraphFormatTool(Tool):
    """Tool to set paragraph format in Word document."""

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "docx_set_paragraph_format"

    @property
    def description(self) -> str:
        return "Set paragraph format (line spacing, indentation, alignment) in a Word document."

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The absolute path to the Word document (.docx)"
                },
                "output_path": {
                    "type": "string",
                    "description": "Output path, omit to overwrite original file"
                },
                "target": {
                    "type": "string",
                    "enum": ["all", "paragraphs", "headings"],
                    "description": "Target scope"
                },
                "line_spacing": {
                    "type": "number",
                    "description": "Line spacing multiplier: 1.0=single, 1.5=1.5x, 2.0=double"
                },
                "line_spacing_pt": {
                    "type": "number",
                    "description": "Fixed line spacing in points (mutually exclusive with line_spacing)"
                },
                "space_before": {
                    "type": "number",
                    "description": "Space before paragraph in points"
                },
                "space_after": {
                    "type": "number",
                    "description": "Space after paragraph in points"
                },
                "first_line_indent": {
                    "type": "number",
                    "description": "First line indent in characters, e.g. 2 means 2-character indent"
                },
                "alignment": {
                    "type": "string",
                    "enum": ["left", "center", "right", "justify"],
                    "description": "Text alignment"
                }
            },
            "required": ["path", "target"]
        }

    async def execute(self, path: str, target: str, output_path: str | None = None,
                      line_spacing: float | None = None, line_spacing_pt: float | None = None,
                      space_before: float | None = None, space_after: float | None = None,
                      first_line_indent: float | None = None, alignment: str | None = None,
                      **kwargs: Any) -> str:
        try:
            file_path = _resolve_path(path, self._workspace, self._allowed_dir)
            if not file_path.exists():
                return f"Error: File not found: {path}"
            if not file_path.suffix.lower() == ".docx":
                return f"Error: Not a .docx file: {path}"

            doc = Document(str(file_path))
            modified_count = 0

            alignment_map = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
            }

            def should_modify_paragraph(para) -> bool:
                if target == "all":
                    return True
                level = _get_heading_level(para)
                if target == "paragraphs":
                    return level is None
                if target == "headings":
                    return level is not None
                return False

            for para in doc.paragraphs:
                if should_modify_paragraph(para):
                    pf = para.paragraph_format
                    
                    if line_spacing is not None:
                        pf.line_spacing = line_spacing
                    if line_spacing_pt is not None:
                        pf.line_spacing = Pt(line_spacing_pt)
                    if space_before is not None:
                        pf.space_before = Pt(space_before)
                    if space_after is not None:
                        pf.space_after = Pt(space_after)
                    if first_line_indent is not None:
                        pf.first_line_indent = Inches(first_line_indent * 0.1389)
                    if alignment and alignment in alignment_map:
                        pf.alignment = alignment_map[alignment]
                    
                    modified_count += 1

            save_path = _resolve_path(output_path, self._workspace, self._allowed_dir) if output_path else file_path
            save_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(save_path))

            return f"Successfully modified {modified_count} paragraphs. Saved to: {save_path}"

        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error setting paragraph format: {str(e)}"


class DocxSetHeadingStyleTool(Tool):
    """Tool to set heading style in Word document."""

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "docx_set_heading_style"

    @property
    def description(self) -> str:
        return "Set style for a specific heading level in a Word document."

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The absolute path to the Word document (.docx)"
                },
                "output_path": {
                    "type": "string",
                    "description": "Output path, omit to overwrite original file"
                },
                "heading_level": {
                    "type": "integer",
                    "enum": [1, 2, 3, 4, 5, 6],
                    "description": "Heading level: 1=Heading 1, 2=Heading 2, etc."
                },
                "font_name": {
                    "type": "string",
                    "description": "Font name"
                },
                "font_size": {
                    "type": "number",
                    "description": "Font size in points"
                },
                "bold": {
                    "type": "boolean",
                    "description": "Whether to make text bold"
                },
                "color": {
                    "type": "string",
                    "description": "Font color in hex"
                },
                "space_before": {
                    "type": "number",
                    "description": "Space before heading in points"
                },
                "space_after": {
                    "type": "number",
                    "description": "Space after heading in points"
                }
            },
            "required": ["path", "heading_level"]
        }

    async def execute(self, path: str, heading_level: int, output_path: str | None = None,
                      font_name: str | None = None, font_size: float | None = None,
                      bold: bool | None = None, color: str | None = None,
                      space_before: float | None = None, space_after: float | None = None,
                      **kwargs: Any) -> str:
        try:
            file_path = _resolve_path(path, self._workspace, self._allowed_dir)
            if not file_path.exists():
                return f"Error: File not found: {path}"
            if not file_path.suffix.lower() == ".docx":
                return f"Error: Not a .docx file: {path}"

            doc = Document(str(file_path))
            modified_count = 0

            for para in doc.paragraphs:
                level = _get_heading_level(para)
                if level == heading_level:
                    for run in para.runs:
                        _set_run_font(run, font_name, font_size, bold, None, color)
                    
                    pf = para.paragraph_format
                    if space_before is not None:
                        pf.space_before = Pt(space_before)
                    if space_after is not None:
                        pf.space_after = Pt(space_after)
                    
                    modified_count += 1

            save_path = _resolve_path(output_path, self._workspace, self._allowed_dir) if output_path else file_path
            save_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(save_path))

            level_name = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六"}.get(heading_level, str(heading_level))
            return f"Successfully modified {modified_count} level-{heading_level} headings. Saved to: {save_path}"

        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error setting heading style: {str(e)}"


class DocxSetTableStyleTool(Tool):
    """Tool to set table style in Word document."""

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "docx_set_table_style"

    @property
    def description(self) -> str:
        return "Set table style (borders, shading, font) in a Word document."

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "The absolute path to the Word document (.docx)"
                },
                "output_path": {
                    "type": "string",
                    "description": "Output path, omit to overwrite original file"
                },
                "table_index": {
                    "type": "integer",
                    "description": "Table index (0-based), -1 for all tables"
                },
                "style_name": {
                    "type": "string",
                    "description": "Built-in style name, e.g. 'Table Grid', 'Light Shading'"
                },
                "font_name": {
                    "type": "string",
                    "description": "Table font name"
                },
                "font_size": {
                    "type": "number",
                    "description": "Table font size in points"
                },
                "header_row": {
                    "type": "boolean",
                    "description": "Whether to set first row as header row"
                },
                "border_style": {
                    "type": "string",
                    "enum": ["none", "single", "double"],
                    "description": "Border style"
                },
                "border_color": {
                    "type": "string",
                    "description": "Border color in hex"
                },
                "cell_shading": {
                    "type": "string",
                    "description": "Cell background color in hex"
                }
            },
            "required": ["path", "table_index"]
        }

    async def execute(self, path: str, table_index: int, output_path: str | None = None,
                      style_name: str | None = None, font_name: str | None = None,
                      font_size: float | None = None, header_row: bool | None = None,
                      border_style: str | None = None, border_color: str | None = None,
                      cell_shading: str | None = None, **kwargs: Any) -> str:
        try:
            file_path = _resolve_path(path, self._workspace, self._allowed_dir)
            if not file_path.exists():
                return f"Error: File not found: {path}"
            if not file_path.suffix.lower() == ".docx":
                return f"Error: Not a .docx file: {path}"

            doc = Document(str(file_path))
            
            if not doc.tables:
                return "Error: No tables found in document"

            tables_to_modify = []
            if table_index == -1:
                tables_to_modify = list(enumerate(doc.tables))
            elif 0 <= table_index < len(doc.tables):
                tables_to_modify = [(table_index, doc.tables[table_index])]
            else:
                return f"Error: Table index {table_index} out of range (0-{len(doc.tables)-1})"

            def set_cell_border(cell, border_type: str, color: str | None):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement("w:tcBorders")
                for border_name in ["top", "left", "bottom", "right"]:
                    border = OxmlElement(f"w:{border_name}")
                    if border_type != "none":
                        border.set(qn("w:val"), border_type)
                        if color:
                            border.set(qn("w:color"), color)
                    else:
                        border.set(qn("w:val"), "nil")
                    tcBorders.append(border)
                tcPr.append(tcBorders)

            def set_cell_shading(cell, color: str):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), color)
                tcPr.append(shd)

            modified_tables = 0
            for idx, table in tables_to_modify:
                if style_name:
                    try:
                        table.style = style_name
                    except Exception:
                        pass

                if header_row and table.rows:
                    first_row = table.rows[0]
                    for cell in first_row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if bold is not None:
                                    run.font.bold = True

                for row in table.rows:
                    for cell in row.cells:
                        if border_style:
                            set_cell_border(cell, border_style, border_color)
                        if cell_shading:
                            set_cell_shading(cell, cell_shading)
                        
                        for para in cell.paragraphs:
                            for run in para.runs:
                                _set_run_font(run, font_name, font_size, None, None, None)

                modified_tables += 1

            save_path = _resolve_path(output_path, self._workspace, self._allowed_dir) if output_path else file_path
            save_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(save_path))

            return f"Successfully modified {modified_tables} table(s). Saved to: {save_path}"

        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error setting table style: {str(e)}"


class DocxFromTemplateTool(Tool):
    """Tool to generate Word document from template."""

    def __init__(self, workspace: Path | None = None, allowed_dir: Path | None = None):
        self._workspace = workspace
        self._allowed_dir = allowed_dir

    @property
    def name(self) -> str:
        return "docx_from_template"

    @property
    def description(self) -> str:
        return "Generate a Word document from a template file with data parameters. Supports Jinja2 syntax for variable substitution."

    @property
    def parameters(self) -> dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "template_path": {
                    "type": "string",
                    "description": "The absolute path to the template file (.docx)"
                },
                "output_path": {
                    "type": "string",
                    "description": "The absolute path for the output file"
                },
                "context": {
                    "type": "object",
                    "description": "Template variables as key-value pairs"
                }
            },
            "required": ["template_path", "output_path", "context"]
        }

    async def execute(self, template_path: str, output_path: str, context: dict[str, Any],
                      **kwargs: Any) -> str:
        try:
            from docxtpl import DocxTemplate

            tpl_path = _resolve_path(template_path, self._workspace, self._allowed_dir)
            if not tpl_path.exists():
                return f"Error: Template file not found: {template_path}"
            if not tpl_path.suffix.lower() == ".docx":
                return f"Error: Not a .docx template: {template_path}"

            out_path = _resolve_path(output_path, self._workspace, self._allowed_dir)
            out_path.parent.mkdir(parents=True, exist_ok=True)

            doc = DocxTemplate(str(tpl_path))
            doc.render(context)
            doc.save(str(out_path))

            return f"Successfully generated document from template. Saved to: {out_path}"

        except ImportError:
            return "Error: docxtpl library not installed. Please install it with: pip install docxtpl"
        except PermissionError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error generating document from template: {str(e)}"
